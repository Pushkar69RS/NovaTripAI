"""Write the Review-1 speaking script as a Word document, from the Markdown.

    uv run python scripts/build_script_docx.py

docs/review1/Travel_Yantra_Phase2_Review1_Script.md is the source; the .docx
beside it is generated, never edited by hand. Headings, paragraphs, bullet
lists, the timing table, quoted cue lines and inline bold/code are carried
over; nothing else is styled, so the document reads like the Markdown does.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs/review1/Travel_Yantra_Phase2_Review1_Script.md"
OUT = ROOT / "docs/review1/Travel_Yantra_Phase2_Review1_Script.docx"

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_runs(paragraph, text: str) -> None:
    """Inline **bold** and `code` into runs; everything else plain."""
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(piece)


def flush(doc, lines: list[str]) -> None:
    if lines:
        para = doc.add_paragraph()
        add_runs(para, " ".join(line.strip() for line in lines))
        lines.clear()


def add_table(doc, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            para = table.cell(r, c).paragraphs[0]
            add_runs(para, cell.strip())
            if r == 0:
                for run in para.runs:
                    run.bold = True


def main() -> int:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11.5)

    pending: list[str] = []
    table: list[list[str]] = []
    for raw in SRC.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("|"):
            flush(doc, pending)
            cells = [c for c in line.strip("|").split("|")]
            if all(set(c.strip()) <= set("-: ") for c in cells):
                continue  # the separator row
            table.append(cells)
            continue
        if table:
            add_table(doc, table)
            table = []
        if not line.strip():
            flush(doc, pending)
            continue
        if line.startswith("# "):
            flush(doc, pending)
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            flush(doc, pending)
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            flush(doc, pending)
            doc.add_heading(line[4:].strip(), level=2)
        elif line.strip() == "---":
            flush(doc, pending)
        elif line.startswith("> "):
            flush(doc, pending)
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Pt(24)
            add_runs(para, line[2:].strip())
            para.runs[0].italic = True if para.runs else None
        elif line.startswith("- "):
            flush(doc, pending)
            para = doc.add_paragraph(style="List Bullet")
            add_runs(para, line[2:].strip())
        elif re.match(r"^\d+\. ", line):
            flush(doc, pending)
            para = doc.add_paragraph(style="List Number")
            add_runs(para, re.sub(r"^\d+\. ", "", line).strip())
        elif line.startswith("[") and line.endswith("]"):
            flush(doc, pending)
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.italic = True
        else:
            pending.append(line)
    flush(doc, pending)
    if table:
        add_table(doc, table)
    for para in doc.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"{OUT.relative_to(ROOT)}  —  {len(doc.paragraphs)} paragraphs")
    return 0


if __name__ == "__main__":
    sys.exit(main())
